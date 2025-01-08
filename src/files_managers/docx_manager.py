import os
from typing import Literal

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches

import constants

from .files_manager import FilesManager

TOrientation = Literal['landscape', 'portrait']


class DocxManager(FilesManager):
    def __init__(self, *, page_url: str) -> None:
        self._page_url = page_url.split('/')[-1]  # seller id
        self._docx_path = self._abs_path(f'assets/{self._page_url}.docx')

        super().__init__(page_url=page_url)

    def create_new_docx(self, path: str | None = None) -> None:
        path = self._abs_path(path if path else f'assets/{self._page_url}.docx')

        doc = Document()
        doc.save(path)
        self._docx_path = path

    def switch_orientation(
        self,
        *,
        orientation: TOrientation,
        docx_path: str | None = None,
        section: int = 0,
    ) -> None:
        docx_path = self._abs_path(docx_path) if docx_path else self._docx_path
        new_orientation: WD_ORIENT = getattr(WD_ORIENT, orientation.upper())

        doc = Document(docx_path)
        doc_section = doc.sections[section]

        if doc_section.orientation != new_orientation:
            doc_section.orientation = new_orientation
            doc_section.page_width = doc_section.page_height
            doc_section.page_height = doc_section.page_width

        doc.save(docx_path)

    def fill_docx_by_dir_pngs(
        self, *, dir_path: str, docx_path: str | None = None
    ) -> None:
        if not self.is_exist(dir_path):
            raise ValueError(f'Directory {dir_path} does not exist')

        get_number = lambda file_name: int(file_name.split('_')[-1].split('.')[0])

        dir_list = sorted(os.listdir(self._abs_path(dir_path)), key=get_number)

        docx_path = self._abs_path(docx_path) if docx_path else self._docx_path
        doc = Document(docx_path)

        for file_name in dir_list:
            if (
                file_name.endswith('.png')
                and str(hash(self._source_page_url)) in file_name
            ):
                doc.add_picture(
                    self._get_photo_path(photo_serial_number=get_number(file_name)),
                    width=Inches(constants.INCHES_IMAGE_COUNT),
                )

        doc.save(docx_path)
