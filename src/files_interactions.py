import os
from typing import Literal

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches
from imagehash import dhash
from PIL import Image

import constants

TOrientation = Literal['landscape', 'portrait']


class FilesManager:
    def __init__(self, *, page_url: str) -> None:
        self._source_page_url = page_url
        self._page_url = page_url.split('/')[-1]  # seller id
        self._docx_path = self._abs_path(f'assets/{self._page_url}.docx')

    @staticmethod
    def _abs_path(file_path: str) -> str:
        return os.path.abspath(file_path)

    def _get_photo_path(self, photo_serial_number: int) -> str:
        return constants.SCREENSHOT_PATH.format(
            hash(self._source_page_url), photo_serial_number
        )

    def delete_file(self, serial_number: int) -> None:
        os.remove(self._get_photo_path(serial_number))

    def create_new_folder(self, path: str) -> None:
        os.mkdir(self._abs_path(path))

    def is_exist(self, path: str) -> bool:
        return os.path.exists(self._abs_path(path))

    def compare_pngs(
        self, first_serial_number: int, second_serial_number: int
    ) -> bool:
        first_path = self._get_photo_path(first_serial_number)
        second_path = self._get_photo_path(second_serial_number)

        return dhash(Image.open(first_path)) == dhash(Image.open(second_path))

    def create_new_docx(self, path: str | None = None) -> None:
        path = self._abs_path(path if path else f'assets/{self._page_url}.docx')

        doc = Document()
        doc.save(path)
        self._docx_path = path

    def switch_orientation(
        self,
        *,
        orientation: TOrientation,
        docx_path: str | None = None,
        section: int = 0,
    ) -> None:
        docx_path = self._abs_path(docx_path) if docx_path else self._docx_path
        new_orientation: WD_ORIENT = getattr(WD_ORIENT, orientation.upper())

        doc = Document(docx_path)
        doc_section = doc.sections[section]

        if doc_section.orientation != new_orientation:
            doc_section.orientation = new_orientation
            doc_section.page_width = doc_section.page_height
            doc_section.page_height = doc_section.page_width

        doc.save(docx_path)

    def fill_docx_by_dir_pngs(
        self, *, dir_path: str, docx_path: str | None = None
    ) -> None:
        if not self.is_exist(dir_path):
            raise ValueError(f'Directory {dir_path} does not exist')

        def get_number(file_name: str) -> int:
            return int(file_name.split('_')[-1].split('.')[0])

        dir_list = sorted(os.listdir(self._abs_path(dir_path)), key=get_number)

        docx_path = self._abs_path(docx_path) if docx_path else self._docx_path
        doc = Document(docx_path)

        for file_name in dir_list:
            if (
                file_name.endswith('.png')
                and str(hash(self._source_page_url)) in file_name
            ):
                doc.add_picture(
                    self._get_photo_path(photo_serial_number=get_number(file_name)),
                    width=Inches(constants.INCHES_IMAGE_COUNT),
                )

        doc.save(docx_path)
